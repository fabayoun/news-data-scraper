import re
from docx.shared import Pt

from document_edits.article_download_state import ArticleDownloadState


def add_article_text(sentences, document, article_download_state):
    """
    Adds article text to document based on input text
    :param sentences: article text
    :param document: document object being passed around
    :return: 
    """

    p = document.add_paragraph()
    # if sentences input is "Text", then article was not downloadable and "Text" is entered as template
    if article_download_state == ArticleDownloadState.FAILED_RESPONSE:
        p_para = p.add_run(sentences)
    # if article downloadable, then find first 3 sentences and add into paragraph
    if article_download_state == ArticleDownloadState.SUCCESS:
        # finds all ". " and turns sentences into list of seperate sentences
        x = re.findall(r'.+', sentences)
        # takes first 3 sentences and merges them
        if len(x) > 3:
            sentences_first_3 = [x[0], x[1], x[2]]
            sentences_merge_first_3 = ". ".join(sentences_first_3)
            # adds sentences to paragraph
            p_para = p.add_run(sentences_merge_first_3)
        else:
            p_para = p.add_run(sentences)
    # formats paragraph
    p_para.font.name = 'Calibri'
    p_para.font.size = Pt(11)
    p_para.italic = True
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    pass
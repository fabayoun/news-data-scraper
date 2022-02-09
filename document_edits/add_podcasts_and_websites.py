from docx.shared import Pt

from document_edits.add_horizontal_break import add_page_break
from document_edits.add_hyperlink import add_hyperlink


energy_news_dict = {
    "Utility Week": r"https://utilityweek.co.uk/",
    "Current-News": r"https://www.current-news.co.uk/",
    "New Power": r"https://www.newpower.info/category/news/",
    "Smart Energy": r"https://www.smart-energy.com/",
    "Green Tech Media": r"https://www.greentechmedia.com/",
    "Networks Online": r"https://networks.online/",
    "Water Briefing": r"https://www.waterbriefing.org/",
    "Water Online": r"https://www.wateronline.com/",
}
energy_podcasts_dict = {
    "GTM: The Energy Gang": r"https://www.greentechmedia.com/podcast/the-energy-gang",
    "GTM: The Interchange": r"https://www.greentechmedia.com/podcast/the-interchange#gs.mJy36kyT",
    "Redefining Energy": r"https://player.fm/series/redefining-energy",
    "Delta Energy Environment": r"https://www.delta-ee.com/TalkingNewEnergy?_cldee=aGFycnkudGF5bG9yQGJhcmluZ2EuY29t&recipientid=contact-945bce7387bce61180ddf4d04bee4450-54ffb43a590b4b7db9e5d77716ddeff1&esid=ca1bc5af-d14f-e911-80e7-ef9585e40074",
    "Talking New Energy": r"https://www.delta-ee.com/TalkingNewEnergy?_cldee=d3RrMzQzQGFvbC5jb20%3d&recipientid=contact-8652e76380a3e71180dffdb8b794821f-7681b9488773437cbea4592f91279093&esid=7b4eaa72-d270-e911-80e7-ef9585e40074",
    "Spotify Talking New Energy": r"https://open.spotify.com/episode/4wqNzmaSozqVjukUgtBsP1?si=3BG654-IQW21-1Otgi2ktw",
}
cga_news_dict = {
    "Civil Service World": r"https://www.civilserviceworld.com/",
    "Institute for Government": r"https://www.instituteforgovernment.org.uk/",
}
cga_podcasts_dict = {
    "Times Red Box": r"https://soundcloud.com/times-comment",
    "BBC Brexitcast": r"https://www.bbc.co.uk/programmes/p05299nl/episodes/player",
    "FT Politics": r"https://www.ft.com/uk-politics-podcast",
    "One Team Gov": r"https://www.oneteamgov.uk/podcast",
}


def add_rn_podcast_and_website_links(document):
    """
    Add R&N podcast and website links
    :param document: document object to be edited
    :return: edited document
    """
    # Add "Energy News: "
    energy_news_header_para = document.add_paragraph()
    energy_news_header = energy_news_header_para.add_run("Energy News: ")
    energy_news_header.font.size = Pt(11)
    energy_news_header.font.name = 'Calibri'
    energy_news_header.bold = True

    for name, url in energy_news_dict.items():
        add_hyperlink(energy_news_header_para, name, url)
        energy_news_header_para.add_run(" // ")

    # Add "Energy Podcasts: "
    energy_podcasts_para = document.add_paragraph()
    energy_podcasts = energy_podcasts_para.add_run("Energy Podcasts: ")
    energy_podcasts.font.size = Pt(11)
    energy_podcasts.font.name = 'Calibri'
    energy_podcasts.bold = True

    # Add Energy Podcasts Links
    for name, url in energy_podcasts_dict.items():
        add_hyperlink(energy_podcasts_para, name, url)
        energy_podcasts_para.add_run(" // ")

    add_page_break(energy_podcasts_para)
    pass


def add_cga_podcast_and_website_links(document):
    """
    Add CGA podcast and website links
    :param document: document object to be edited
    :return: edited document
    """
    # Add "CGA News: "
    cga_news_header_para = document.add_paragraph()
    cga_news_header = cga_news_header_para.add_run("CGA News: ")
    cga_news_header.font.size = Pt(11)
    cga_news_header.font.name = 'Calibri'
    cga_news_header.bold = True

    # Add CGA News links
    for name, url in cga_news_dict.items():
        add_hyperlink(cga_news_header_para, name, url)
        cga_news_header_para.add_run(" // ")

    # Add "CGA Podcasts: "
    cga_podcast_header_para = document.add_paragraph()
    cga_podcast_header = cga_podcast_header_para.add_run("CGA Podcasts: ")
    cga_podcast_header.font.size = Pt(11)
    cga_podcast_header.font.name = 'Calibri'
    cga_podcast_header.bold = True

    # Add CGA Podcast links
    for name, url in cga_podcasts_dict.items():
        add_hyperlink(cga_podcast_header_para, name, url)
        cga_podcast_header_para.add_run(" // ")

    add_page_break(cga_podcast_header_para)
    pass

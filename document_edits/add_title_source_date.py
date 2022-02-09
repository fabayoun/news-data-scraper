import re
from docx.shared import Pt, RGBColor

##dictionary which defines source, based on url. Useful for creating formatted source
dict_source = {
    "current-news": "Current News",
    "thesun": "The Sun",
    "bbc": "BBC",
    "goodenergy": "Good Energy",
    "bnef": "BNEF",
    "instituteforgovernment": "Institute for Government",
    "civilserviceworld": "CSW",
    "theguardian": "The Guardian",
    "itv": "ITV",
    "thedrum": "The Drum",
    "independent": "The Independent",
    "Source": "Source", #for template when article can't be downloaded (currently not in use)
    "sky": "Sky News",
    "thetimes": "The Times",
    "ft": "The Financial Times",
    "moneysavingexpert": "The Money Saving Expert",
    "telegraph": "The Telegraph",
    "theenergyst": "The Energyst",
    "engerati": "Engerati",
    "openaccessgovernment": "Open Access Government",
    "publictechnology": "Public Technology",
    "room151": "Room 515",
    "consultancy": "Consultancy.UK",
    "wired": "Wired",
    "itproportal": "ITProPortal",
    "tortoisemedia": "Tortoise Media",
    "zap-map": "Zap Map",
    "cityam": "City A.M.",
    "utilityweek": "Utility Week",
    "arstechnica": "Arstechnica",
    "bpchargemaster": "Bpchargemaster",
    "weforum": "Weforum",
    "politicshome": "Politicshome",
    "energyvoice": "Energy Voice",
    "ofwat": "Ofwat",
    "forbes": "Forbes",
    "theregister": "The Register"
}

def add_title_source_date(title, source, date, document):
    """Add formatted title, source and date to document
    :param title: title added as given
    :param source: source determined from URL and library, if not in library states: "Add Source"
    :param date: date as given, if doesn't exist states: "No Date"
    :return:
    """
    ##defines t which adds title paragraph where title/source/date are added to
    t = document.add_paragraph()

    ##add title
    #___adds title text "title" and formats
    t_title = t.add_run(title)
    t_title.font.size = Pt(13)
    t_title.font.name = 'Calibri'
    t_title.bold = True

    ##add source to document based on "url" and "dict_source" dictionary
    #___splits url into list and takes useful source snippet by identifying list item before "co", "com", or "org"
    #___e.g takes XXX in www.XXX.com

    #decide whether input is url or if it is inputted text (e.g for case that string couldn't be downloaded)
    if "http" in source:
        #split into  list
        source_split = re.split("[/.]", source)
        #identify list item before "co", "com", "org" or "net"
        if "co" in source_split:
            n = source_split.index("co")
            check_source = source_split[n-1]
        elif "com" in source_split:
            n = source_split.index("com")
            check_source = source_split[n-1]
        elif "org" in source_split:
            n = source_split.index("org")
            check_source = source_split[n-1]
        elif "net" in source_split:
            n = source_split.index("net")
            check_source = source_split[n-1]
        else:
            check_source = source
    else:
        check_source = source

    #checks if url exists in source dictionary
    if check_source in dict_source:
        #if url exists in dictionary, add source
        url_s = dict_source[check_source]
        t_source = t.add_run(' (' + url_s)
        t_source.font.name = 'Calibri'
        t_source.font.size = Pt(11)
        t_source.bold = True
        t_source.italic = True
    else:
        #when source doesn't exist in dictionary this formats text to state "Add Source" in red.
        t_source = t.add_run(' (' + "Add Source to Dictionary")
        t_source.font.name = 'Calibri'
        t_source.font.size = Pt(15)
        t_source.bold = True
        t_source.italic = True
        t_source.font.color.rgb = RGBColor(255,0,0)

    ##add date to document if it exists
    #creates object to check date type
    date_time_obj = date
    #checks if date exists and if date doesn't exist, No date is added
    if date_time_obj == None:
        t_date = t.add_run(' - No Date)')
    #checks if date is named "Date" as this only is the case when article failed to download and template is required. Adds "Date" if so.
    elif date_time_obj == "Date":
        t_date = t.add_run(' - ' + date_time_obj + ')')
    #if date exists, timestamp is removed and date is turned into formatted string
    else:
        date_only = date_time_obj.date()
        date_str = date_only.strftime('%d/%m/%Y')
        t_date = t.add_run(' - ' + date_str + ')')

    #final formatting of date
    t_date.font.name = 'Calibri'
    t_date.font.size = Pt(11)
    t_date.bold = True
    t_date.italic = True


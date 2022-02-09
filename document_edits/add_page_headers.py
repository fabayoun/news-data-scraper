from docx.shared import Pt

###Add CGA Heading and Author
def add_CGA_heading_author(document):
    ##dd space before Heading
    space = document.add_paragraph()

    ##Add CGA Heading
    h_CGA = document.add_paragraph()
    h_CGA_h = h_CGA.add_run("Central Governments Advisory")
    h_CGA_h.font.size = Pt(15)
    h_CGA_h.font.name = 'Calibri'
    h_CGA_h.bold = True

    ##Add CGA Author
    h_CGA_a = document.add_paragraph()
    h_CGA_a_a = h_CGA_a.add_run("By Harry Allport & Rhiannon Evans")
    h_CGA_a_a.font.size = Pt(11)
    h_CGA_a_a.font.name = 'Calibri'

    ##Add CGA Heading
    space = document.add_paragraph()


def add_RN_heading(document):
    h_RN = document.add_paragraph()
    h_RN_h = h_RN.add_run("Retail & Networks")
    h_RN_h.font.size = Pt(15)
    h_RN_h.font.name = 'Calibri'
    h_RN_h.bold = True
    space = document.add_paragraph()
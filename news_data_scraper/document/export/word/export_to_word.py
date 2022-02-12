import logging
from pathlib import Path

from docx import Document

from news_data_scraper.document.export.word.add_page_headers import add_heading_and_author
from news_data_scraper.document.export.word.add_title_source_date import add_article
from news_data_scraper.document.export.word.add_podcasts_and_websites import add_podcast_and_website_links
from news_data_scraper.document.document_classes import NewsDocument
from news_data_scraper.scraper.bu_tags import BuTag


def export_to_word(document_contents: NewsDocument, output_file_name: str, output_file_path: Path) -> None:
    document = Document()
    format_document(document, document_contents)
    document.save(output_file_path)
    logging.info(f"Completed: {output_file_name}")


def format_document(document: Document, document_contents: NewsDocument) -> None:
    for section in document_contents.business_unit_sections:
        add_heading_and_author(document, section)
        add_article(document, section)
        if section.bu_tag != BuTag.NONE:
            add_podcast_and_website_links(document, section)
        document.add_paragraph()

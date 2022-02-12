from docx import Document
from docx.shared import Pt

from news_data_scraper.document.export.word.add_horizontal_break import add_page_break
from news_data_scraper.document.export.word.add_hyperlink import add_hyperlink
from news_data_scraper.document.document_feature_classes import BusinessUnitSection


def add_podcast_and_website_links(document: Document, section: BusinessUnitSection) -> None:
    """
    Add R&N podcast and website links
    :param document: document object to be edited
    :param section: news section which we are adding podcast and website links for
    :return: edited document
    """
    news_header_para = document.add_paragraph()
    news_header = news_header_para.add_run("News Sites: ")
    news_header.font.size = Pt(11)
    news_header.font.name = 'Calibri'
    news_header.bold = True

    for name, url in section.all_news_websites.news_sites.items():
        add_hyperlink(news_header_para, name, url)
        news_header_para.add_run(" // ")

    # Add "Energy Podcasts: "
    podcasts_para = document.add_paragraph()
    podcasts = podcasts_para.add_run("Podcasts: ")
    podcasts.font.size = Pt(11)
    podcasts.font.name = 'Calibri'
    podcasts.bold = True

    # Add Energy Podcasts Links
    for name, url in section.all_news_websites.news_sites.items():
        add_hyperlink(podcasts_para, name, url)
        podcasts_para.add_run(" // ")

    add_page_break(podcasts_para)
    pass


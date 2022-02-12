from docx.shared import Pt
from docx import Document

from news_data_scraper.document.document_classes import BusinessUnitSection


def add_article(document: Document, section: BusinessUnitSection) -> None:
    """Add formatted title, source and date to document"
    :param document: document object
    :param section: the section (R&N/CGA to be formatted)
    :return:
    """

    for article in section.all_news_articles.articles:
        t = document.add_paragraph()

        t_title = t.add_run(article.title)
        t_title.font.size = Pt(13)
        t_title.font.name = 'Calibri'
        t_title.bold = True

        t_source = t.add_run(f" ({article.source}")
        t_source.font.name = 'Calibri'
        t_source.font.size = Pt(11)
        t_source.bold = True
        t_source.italic = True

        t_date = t.add_run(f" - {article.date})")
        t_date.font.name = 'Calibri'
        t_date.font.size = Pt(11)
        t_date.bold = True
        t_date.italic = True

        p = document.add_paragraph()
        p_para = p.add_run(article.text)

        p_para.font.name = 'Calibri'
        p_para.font.size = Pt(11)
        p_para.italic = True
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        document.add_paragraph()

    pass








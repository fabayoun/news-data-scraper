from docx.shared import Pt
from docx import Document

from news_data_scraper.other_document_features.document_feature_classes import NewsDocument, NewsSection


def add_heading_and_author(document: Document, section: NewsSection):
    """
    Add CGA Heading and Author
    :param document: document object to be edited
    :return: edited document
    """
    document.add_paragraph()

    # Add BU Header
    header_cga = document.add_paragraph().add_run(section.business_unit_header.header)
    header_cga.font.size = Pt(15)
    header_cga.font.name = 'Calibri'
    header_cga.bold = True

    # Add CGA Author
    author_cga = document.add_paragraph().add_run(section.business_unit_header.author)
    author_cga.font.size = Pt(11)
    author_cga.font.name = 'Calibri'

    document.add_paragraph()

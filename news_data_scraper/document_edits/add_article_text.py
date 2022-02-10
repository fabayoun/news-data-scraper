import re
from docx.shared import Pt

from news_data_scraper.document_edits.article_download_state import ArticleDownloadState


def add_article_text(sentences, document, article_download_state):
    """
    Adds article text to document based on input text
    :param sentences: article text
    :param document: document object being passed around
    :return: 
    """

    p = document.add_paragraph()

    if article_download_state == ArticleDownloadState.FAILED_RESPONSE:
        p_para = p.add_run(sentences)

    if article_download_state == ArticleDownloadState.SUCCESS:
        x = re.findall(r'.+', sentences)
        if len(x) > 3:
            sentences_first_3 = [x[0], x[1], x[2]]
            sentences_merge_first_3 = ". ".join(sentences_first_3)
            p_para = p.add_run(sentences_merge_first_3)
        else:
            p_para = p.add_run(sentences)

    p_para.font.name = 'Calibri'
    p_para.font.size = Pt(11)
    p_para.italic = True
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    pass
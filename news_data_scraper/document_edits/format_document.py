from docx import Document

from news_data_scraper.document_edits.add_horizontal_break import add_page_break
from news_data_scraper.document_edits.add_page_headers import add_heading_and_author
from news_data_scraper.document_edits.add_title_source_date import add_article
from news_data_scraper.document_edits.add_podcasts_and_websites import add_podcast_and_website_links
from news_data_scraper.other_document_features.document_feature_classes import NewsDocument
from news_data_scraper.scraper.bu_tags import BuTag


def format_document(document: Document, document_contents: NewsDocument) -> None:
    for section in document_contents.business_unit_sections:
        add_heading_and_author(document, section)
        add_article(document, section)
        if section.bu_tag != BuTag.NONE:
            add_podcast_and_website_links(document, section)
        document.add_paragraph()




